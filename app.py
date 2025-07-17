# app.py — com melhorias finais refinadas para Render

import os, re, time, json, shutil, random, asyncio, unicodedata
from urllib.parse import urljoin

import nest_asyncio
nest_asyncio.apply()

import requests
from bs4 import BeautifulSoup
from docx import Document
import gradio as gr
from openai import AsyncOpenAI
import re
import difflib

HEADERS = {
    'User-Agent': (
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
        'AppleWebKit/537.36 (KHTML, like Gecko) '
        'Chrome/91.0.4472.124 Safari/537.36'
    )
}

# --- Funções auxiliares ---
def tempo_espera(MIN_TIME=5, MAX_TIME=9, contexto="esperando..."):
    frases = [
        "⌛ Fazendo uma pausa pra não fritar os servidores...",
    ]
    tempo = random.uniform(MIN_TIME, MAX_TIME)
    print(f"\n{random.choice(frases)} ({contexto})")
    time.sleep(tempo)
    print(f"...retomando após {tempo:.2f} segundos.\n")

def is_valid_url(url):
    return url.startswith("http") and ".html" in url

def is_irrelevant_text(texto):
    termos_ruins = [
        "solicite um orçamento", "contate-nos", "compartilhe:", "siga-nos",
        "nossa equipe", "serviço ao cliente", "entre em contato",
        "política de privacidade", "fale conosco", "acesse nosso site",
        "atendimento ao cliente", "cnpj:", "telefone:", "email:"
    ]
    texto_lower = texto.lower().strip()
    return any(t in texto_lower for t in termos_ruins)

def coletar_links_artigos(pagina_url):
    try:
        response = requests.get(pagina_url, headers=HEADERS)
        response.raise_for_status()
    except Exception as e:
        print(f"[Erro] Falha ao acessar {pagina_url}: {e}")
        return []

    sopa = BeautifulSoup(response.text, 'html.parser')

    for seletor in ['footer', '.footer', '#footer', '.site-footer', '.rodape', '.legal', '.copyright']:
        for el in sopa.select(seletor):
            el.decompose()

    keywords = [
        "demonstração de produto", "solicite um orçamento", "representantes profissionais",
        "valorizamos o seu negócio", "todas as marcas", "logotipos da Tennant", "propriedade da Tennant"
    ]
    paragraphs = sopa.find_all(['p', 'div', 'span'])
    for p in paragraphs:
        if any(k in p.get_text(strip=True).lower() for k in keywords):
            p.decompose()

    todos_a = sopa.find_all('a', href=True, title=True)
    links = []

    for a in todos_a:
        href = a['href']
        title = a['title'] or a.text.strip()

        if not href.startswith('http'):
            href = urljoin(URL_BASE, href)

        if any(excl in href for excl in ['cart', 'contact', 'solicitud', 'linkedin', 'facebook', 'twitter']):
            continue

        if not is_valid_url(href):
            continue

        if '/blog/' not in href and PAIS not in ['ja_jp', 'zh_cn', 'ko_kr']:
            continue

        links.append({'title': title.strip(), 'href': href})

    vistos = set()
    links_unicos = []
    for l in links:
        if l['href'] not in vistos:
            vistos.add(l['href'])
            links_unicos.append(l)

    print(f"🔗 {len(links_unicos)} links válidos extraídos da página.")
    return links_unicos

def get_article_content(article_url):
    try:
        tempo_espera(7.5, 9.5, contexto="esperando antes de coletar o artigo")
        response = requests.get(article_url, headers=HEADERS, timeout=15)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, 'html.parser')

        # Remover seções irrelevantes
        for seletor in ['nav', '.nav', '#nav', '.breadcrumbs', '.category-list']:
            for el in soup.select(seletor):
                el.decompose()

        title_tag = soup.find('h1')
        title = title_tag.get_text(strip=True) if title_tag else "Título não encontrado"

        paragraphs = soup.find_all(['p', 'h2', 'h3'])
        raw_content = [para.get_text() for para in paragraphs if para.get_text(strip=True)]
        raw_content = [c.strip() for c in raw_content if c.strip()]

        # Filtro de conteúdo irrelevante
        filtered = []
        seen = set()
        for c in raw_content:
            c_clean = c.strip()
            if len(c_clean) > 40 and not is_irrelevant_text(c_clean) and c_clean.lower() not in seen:
                filtered.append(c_clean)
                seen.add(c_clean.lower())

        return title, filtered

    except Exception as e:
        print(f"[Erro Geral] {e}")
        return None, []

def clean_filename(s):
    proibidos = '<>:"/\\|?*'
    for char in proibidos:
        s = s.replace(char, '')
    s = s.strip()
    if len(s) > 50:
        s = s[:50]
    return s.replace(' ', '_')

def limpar_xml(texto):
    return ''.join(
        c for c in texto
        if c == '\n' or c == '\r' or c == '\t' or 32 <= ord(c) <= 126 or ord(c) >= 160
    )

def salvar_conteudo(title, content, pasta):
    nome_arquivo = clean_filename(title)
    path_docx = os.path.join(pasta, f"{nome_arquivo}.docx")
    doc = Document()
    doc.add_heading(limpar_xml(title), level=1)
    for p in content:
        doc.add_paragraph(f"• {p}")
    doc.save(path_docx)
    print(f"💾 Documento salvo: {path_docx}")

# === Config ===
print("🔧 Iniciando app.py...")

with open("/etc/secrets/OPENAI_KEY") as f:
    api_key = f.read().strip()
    print("🔑 OPENAI_KEY carregada:", api_key[:6], "... (ocultado)")
if not api_key:
    raise ValueError("⚠️ Variável de ambiente 'OPENAI_KEY' não encontrada.")
client = AsyncOpenAI(api_key=api_key)

with open("/etc/secrets/SCRAPER_PASSWORD") as f:
    SENHA = f.read().strip()
    print("🔐 SCRAPER_PASSWORD carregada com sucesso.")
if not SENHA:
    raise ValueError("⚠️ Variável de ambiente 'SCRAPER_PASSWORD' não encontrada.")

URL_BASE = "https://www.tennantco.com"
MAPA_PAISES = {
    'pt_br': 'Brasil', 'en_us': 'Estados Unidos', 'en_ca': 'Canadá',
    'en_au': 'Austrália e nova zelandia', 'en_za': 'África do Sul', 'en_gb': 'Reino Unido',
    'es_es': 'Espanha', 'es_mx': 'México', 'fr_fr': 'França', 'nl_nl': 'Holanda',
    'en_eu': 'Europa (outros paises)', 'en_ap': 'Ásia (outros paises)',
    'en_la': 'america latina (outros paises)', 'es_la': 'america latina (outros paises)',
    'de_de': 'Alemanha', 'it_it': 'Itália', 'ja_jp': 'Japão', 'zh_cn': 'China', 'pt_pt': 'Portugal'
}

def normalizar(texto):
    return unicodedata.normalize('NFKD', texto).encode('ASCII', 'ignore').decode('ASCII').lower().strip()

def limpar_pasta_resultados(caminho):
    if os.path.exists(caminho):
        shutil.rmtree(caminho)
    os.makedirs(caminho)

MAPA_NOMES = {normalizar(nome): codigo for codigo, nome in MAPA_PAISES.items()}
ALIASES_PAISES = {
    "canguru": "en_au", "boomerang": "en_au", "sidney": "en_au", "aussie": "en_au", "kiwi": "en_au",
    "samba": "pt_br", "carnaval": "pt_br", "taco": "es_mx", "mariachi": "es_mx",
    "eiffel": "fr_fr", "croissant": "fr_fr", "molde": "nl_nl", "tulipa": "nl_nl",
    "shinkansen": "ja_jp", "samurai": "ja_jp", "dragao": "zh_cn", "mao": "zh_cn",
    "realeza": "en_gb", "londres": "en_gb", "snow": "en_ca", "hockey": "en_ca",
    "bavaria": "de_de", "oktoberfest": "de_de", 'RJ': 'pt_br', 'SP': 'pt_br'
}

# === Tradução GPT refinada ===
async def traduzir_e_formatar_gpt(textos, destino='português Brasil'):
    resultados = []
    modelo = "gpt-4o-mini"
    def agrupar_paragrafos(paragrafos, max_chars=800):
    """
    Agrupa os parágrafos em blocos de tamanho definido, dando preferência por títulos <h2>.
    Reduz o número de tokens e evita blocos excessivamente longos.
    """
    blocos = []
    buffer = ""
    
    for par in paragrafos:
        # Se o buffer atual mais o próximo parágrafo exceder o tamanho, salva o buffer atual.
        if len(buffer) + len(par) + 1 <= max_chars:
            buffer += par + "\n"
        else:
            if buffer.strip():
                blocos.append(buffer.strip())
            buffer = par + "\n"  # Começa novo bloco

    if buffer.strip():
        blocos.append(buffer.strip())

    return blocos
    total_prompt_tokens = 0
    total_completion_tokens = 0

    termos_remover = [
        "entre em contato", "descubra como", "solicite um orçamento", "fale conosco",
        "nossa equipe de vendas", "está pronta para ajudar", "demonstração de produto",
        "visite nosso site", "nos siga nas redes", "cnpj:", "telefone:", "email:", "diretores da empresa"
    ]

    termos_preservar = ["BrainOS", "ec-H2O NanoClean", "T16AMR", "i-mop", "LiDAR", "3D", "AMR"]

    textos_corrigidos = []
    for texto in textos:
        for termo in termos_preservar:
            texto = texto.replace(termo, f"{termo}")
        textos_corrigidos.append(texto)

    # Pré-filtragem por relevância
    textos_filtrados = []
    for t in textos_corrigidos:
        if not any(termo in t.lower() for termo in termos_remover):
            textos_filtrados.append(t)

    # Agrupamento de blocos com prevenção de duplicatas e curtos
    vistos_hash = set()
    for texto in textos_filtrados:
        texto_limpo = texto.strip()
        if not texto_limpo or len(texto_limpo.split()) < 5:
            continue

        hash_bloco = hash(re.sub(r'\W+', '', texto_limpo.lower()))
        if hash_bloco in vistos_hash:
            continue
        vistos_hash.add(hash_bloco)

        if len(buffer) + len(texto_limpo) + 1 < max_chars:
            buffer += " " + texto_limpo
        else:
            blocos.append(buffer.strip())
            buffer = texto_limpo
    if buffer:
        blocos.append(buffer.strip())

    # Tradução de cada bloco
    for bloco in blocos:
        system_msg = {
            "role": "system",
            "content": (
                "Você é um tradutor profissional. Traduza para o português do Brasil com fidelidade, coesão, fluidez e tom editorial.\n\n"
                "Regras:\n"
                "1) Não adicione chamadas promocionais ou institucionais.\n"
                "2) Preserve nomes técnicos e marcas (ex: i-mop, ec-H2O, CS5, T500).\n"
                "3) Ignore rodapés, categorias e menus.\n"
                "4) Use “esfregão” para mop e “lavadora de pisos” ou “esfregadora” para scrubber.\n"
                "5) Evite repetições e frases soltas; traduza com naturalidade.\n"
                "6) Não marque os termos preservados.\n"
                "7) Se o conteúdo estiver vazio, irrelevante ou genérico, ignore sem responder com uma mensagem padrão. Apenas não gere nada."
            )
        }
        user_msg = {"role": "user", "content": bloco}

        try:
            resposta = await client.chat.completions.create(
                model=modelo,
                messages=[system_msg, user_msg],
                temperature=0.3,
                max_tokens=2000,
                n=1
            )

            texto_traduzido = resposta.choices[0].message.content.strip()
            usage = resposta.usage
            prompt_tokens = usage.prompt_tokens
            completion_tokens = usage.completion_tokens
            total_tokens = usage.total_tokens

            print(f"\U0001F4CA Tokens usados neste bloco → Prompt: {prompt_tokens}, Resposta: {completion_tokens}, Total: {total_tokens}")

            total_prompt_tokens += prompt_tokens
            total_completion_tokens += completion_tokens

            paragrafos = []
            linhas_vistas = set()
            for linha in texto_traduzido.split('\n'):
                linha_limpa = linha.strip()
                if not linha_limpa:
                    continue
                if linha_limpa.lower() in linhas_vistas:
                    continue
                linhas_vistas.add(linha_limpa.lower())
                paragrafos.append(linha_limpa)

            resultados.extend(paragrafos)
            await asyncio.sleep(random.uniform(1.2, 2.0))

        except Exception as e:
            print(f"[Erro GPT Tradução] {e}")
            resultados.append(bloco)

    print(f"\n\U0001F4C8 Token log total:")
    print(f"  🔹 Prompt tokens: {total_prompt_tokens}")
    print(f"  🔹 Completion tokens: {total_completion_tokens}")
    print(f"  🔹 Tokens totais: {total_prompt_tokens + total_completion_tokens}\n")

    # 🔄 Filtro final para remover seções duplicadas com mesmo título (ex: "1. Eficiência...", "2. ...")
    secoes_vistas = set()
    secoes_filtradas = []
    for paragrafo in resultados:
        titulo_match = re.match(r'^(\d\.\s?.+?)(?=[\:\.\-–])', paragrafo)
        if titulo_match:
            titulo_normalizado = titulo_match.group(1).strip().lower()
            if titulo_normalizado in secoes_vistas:
                continue  # ignora duplicata
            secoes_vistas.add(titulo_normalizado)
        secoes_filtradas.append(paragrafo)

    # Remoção de parágrafos altamente semelhantes
    final_resultados = []
    for p in paragrafos:
        p_normalizado = p.strip().lower()
        if not any(difflib.SequenceMatcher(None, p_normalizado, existente.lower()).ratio() > 0.85 for existente in final_resultados):
            final_resultados.append(p)

    # Checagem mínima de tópicos
    num_topicos = sum(1 for p in final_resultados if re.match(r'^\d\.', p))
    if num_topicos < 5:
        print(f"⚠️ Alerta: apenas {num_topicos}/5 tópicos detectados na tradução.")

    return final_resultados, {
        "prompt_tokens": total_prompt_tokens,
        "completion_tokens": total_completion_tokens,
        "total_tokens": total_prompt_tokens + total_completion_tokens
    }
async def run(pais, alias, qtd_artigos):
    global PAIS
    entrada = normalizar(pais)
    if entrada in MAPA_PAISES:
        codigo = entrada
    elif entrada in MAPA_NOMES:
        codigo = MAPA_NOMES[entrada]
    elif entrada in ALIASES_PAISES:
        codigo = ALIASES_PAISES[entrada]
    elif alias.strip() in MAPA_PAISES:
        codigo = alias.strip()
    else:
        return f"❌ País inválido: {pais}", [], gr.update(visible=True)

    PAIS = codigo  # <- aqui, corretament

    nome_pais = MAPA_PAISES[codigo]
    url_blog = f"{URL_BASE}/{codigo}/blog.html"
    pasta = os.path.abspath(f"resultados/{nome_pais}")
    limpar_pasta_resultados(pasta)

    links = coletar_links_artigos(url_blog)

    vistos_titulos = set()
    vistos_hashes = set()
    processados = 0
    falhas = []
    token_totais = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}

    for art in links:
        if processados >= int(qtd_artigos):
            break

        title, content = get_article_content(art['href'])
        if not content:
            falhas.append(title)
            continue

        titulo_normalizado = normalizar(title)
        hash_artigo = hash(" ".join(content).strip().lower())
        if titulo_normalizado in vistos_titulos or hash_artigo in vistos_hashes:
            continue

        vistos_titulos.add(titulo_normalizado)
        vistos_hashes.add(hash_artigo)

        traduzido, token_log = await traduzir_e_formatar_gpt(content)
        salvar_conteudo(title, traduzido, pasta)
        processados += 1

        for k in token_totais:
            token_totais[k] += token_log.get(k, 0)

    arquivos_docx = sorted([os.path.join(pasta, f) for f in os.listdir(pasta) if f.endswith(".docx")])

    if falhas:
        print('easter egg')
        for f in falhas:
            print(f" - {f}")

    resumo_tokens = (
        f"\n📊 Uso de tokens:"
        f"\n• Prompt: {token_totais['prompt_tokens']}"
        f"\n• Resposta: {token_totais['completion_tokens']}"
        f"\n• Total: {token_totais['total_tokens']}"
    )

    return "✅ Concluído!" + resumo_tokens, arquivos_docx, gr.update(visible=False)

def checar_senha(senha_input):
    return (gr.update(visible=False), gr.update(visible=True)) if senha_input == SENHA else (gr.update(visible=True), gr.update(visible=False))

# === Interface ===
with gr.Blocks(title="Tennant Translator") as demo:
    gr.HTML("""<style>#arquivos_box .wrap.svelte-1ipelgc { max-height: 300px; overflow-y: auto; display: flex; flex-direction: column; gap: 4px; }.gradio-container { max-width: 95vw !important; margin: auto !important; }.gr-textbox, .gr-number { flex: 1 1 48%; min-width: 300px; }.gr-button { width: 100%; }</style>""")

    with gr.Row(visible=True) as login_box:
        senha_input = gr.Textbox(label="Digite a senha", type="password")
        btn_login = gr.Button("Entrar")

    with gr.Row(visible=False) as app_box:
        gr.Markdown("## 🌐 Tradutor de Artigos Tennant (GPT-4o-mini)")
        with gr.Row():
            pais = gr.Textbox(label="País ou termo", placeholder="ex: taco, samba, holanda")
            alias = gr.Textbox(label="Alias (se for novo)", placeholder="ex: es_mx")
            qtd = gr.Number(label="Qtd. artigos", value=3)

        status = gr.Textbox(label="Status do processo", interactive=False)
        with gr.Column(elem_id="arquivos_box"):
            arquivos = gr.File(label="Arquivos traduzidos (.docx)", file_types=[".docx"], file_count="multiple")

        btn = gr.Button("Iniciar")
        btn.click(run, inputs=[pais, alias, qtd], outputs=[status, arquivos, alias])

    btn_login.click(checar_senha, inputs=senha_input, outputs=[login_box, app_box])

if __name__ == "__main__":
    import sys
    port = int(os.getenv("PORT", 7860))
    demo.launch(server_name="0.0.0.0", server_port=port)
