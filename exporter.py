import os
import requests
import mimetypes
from urllib.parse import urlparse
from docx import Document
from docx.shared import Inches, RGBColor
from PIL import Image
from utils import clean_filename, limpar_xml

HEADERS = {
    'User-Agent': (
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
        'AppleWebKit/537.36 (KHTML, like Gecko) '
        'Chrome/91.0.4472.124 Safari/537.36'
    )
}

DESTAQUES = [
    "Um dos maiores receios",
    "A expectativa é que",
    "Eficiência por meio da colaboração",
    "3. Comprovação da limpeza",
    "4. Capacidade dos funcionários",
    "5. Uso otimizado de recursos"
]

def aplicar_destaque_manual(paragraph, texto):
    for frase in DESTAQUES:
        if texto.startswith(frase):
            run = paragraph.add_run(texto)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 204)  # Azul
            return True
    return False

def baixar_imagem(url, pasta_destino):
    try:
        print(f"\n🔽 Tentando baixar imagem: {url}", flush=True)
        response = requests.get(url, headers=HEADERS, stream=True, timeout=10)
        response.raise_for_status()

        content_type = response.headers.get('Content-Type')
        if not content_type or not content_type.lower().startswith("image/"):
            print(f"[⚠️ Tipo de conteúdo inválido ou ausente] {url}")
            return None

        ext = mimetypes.guess_extension(content_type.lower()) or '.jpg'
        print(f"📦 Tipo de conteúdo: {content_type} | Extensão detectada: {ext}", flush=True)

        nome_url = os.path.basename(urlparse(url).path).split("?")[0]
        if not nome_url:
            nome_url = f"img_{hash(url)}"
        if not os.path.splitext(nome_url)[1]:
            nome_url += ext

        nome_arquivo = clean_filename(nome_url)
        caminho = os.path.join(pasta_destino, nome_arquivo)

        with open(caminho, 'wb') as f:
            f.write(response.content)

        # Verifica se a imagem é válida com Pillow
        try:
            with Image.open(caminho) as img:
                img.verify()
        except Exception as e:
            print(f"[⚠️ Imagem corrompida ou inválida] {url}: {e}")
            return None

        print(f"✅ Imagem salva em: {caminho}", flush=True)
        return caminho

    except Exception as e:
        print(f"[❌ Erro ao baixar imagem] {url}: {e}", flush=True)
        return None

def reparar_imagem(caminho_original):
    try:
        with Image.open(caminho_original) as img:
            rgb = img.convert('RGB')
            caminho_corrigido = caminho_original.replace(".", "_reparada.", 1)
            rgb.save(caminho_corrigido, format="JPEG")
            print(f"[🛠️ Imagem reparada e salva como] {caminho_corrigido}", flush=True)
            return caminho_corrigido
    except Exception as e:
        print(f"[⚠️ Erro ao reparar imagem] {caminho_original}: {e}", flush=True)
        return None

def salvar_conteudo_em_docx(titulo, elementos, pasta_saida, url_origem=None):
    """Salva o conteúdo traduzido em um arquivo .docx com texto e imagens"""
    nome_arquivo = clean_filename(titulo)
    caminho = os.path.join(pasta_saida, f"{nome_arquivo}.docx")
    os.makedirs(pasta_saida, exist_ok=True)

    doc = Document()

    if url_origem:
        doc.add_paragraph(f"🔗 Artigo original: {url_origem}", style="Intense Quote")

    doc.add_heading(limpar_xml(titulo), level=1)

    print(f"\n📝 Iniciando documento: {nome_arquivo}", flush=True)
    total_imgs = 0
    total_paragrafos = 0

    for item in elementos:
        tipo = item['tipo']
        conteudo = item['conteudo'].strip()

        if tipo == 'h2':
            doc.add_paragraph(conteudo, style='Heading 2')
            print(f"🔹 H2: {conteudo[:50]}...", flush=True)

        elif tipo == 'h3':
            doc.add_paragraph(conteudo, style='Heading 3')
            print(f"🔸 H3: {conteudo[:50]}...", flush=True)

        elif tipo == 'p':
            paragraph = doc.add_paragraph()
            if not aplicar_destaque_manual(paragraph, conteudo):
                paragraph.add_run(conteudo)
            total_paragrafos += 1

        elif tipo == 'img':
            img_path = baixar_imagem(conteudo, pasta_saida)
            if img_path:
                try:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    total_imgs += 1
                    print(f"🖼️ Imagem inserida: {img_path}", flush=True)
                except Exception as e1:
                    print(f"[⚠️ Falha ao inserir imagem original] {img_path}: {e1}", flush=True)
                    img_corrigida = reparar_imagem(img_path)
                    if img_corrigida:
                        try:
                            paragraph = doc.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_corrigida, width=Inches(5.5))
                            total_imgs += 1
                            print(f"🖼️ Imagem reparada inserida: {img_corrigida}", flush=True)
                        except Exception as e2:
                            print(f"[❌ Erro ao inserir imagem reparada] {img_corrigida}: {e2}", flush=True)

    doc.save(caminho)
    print(f"\n💾 Arquivo salvo com sucesso: {caminho}", flush=True)
    print(f"📊 Estatísticas: {total_paragrafos} parágrafos | {total_imgs} imagens\n", flush=True)
    return caminho
